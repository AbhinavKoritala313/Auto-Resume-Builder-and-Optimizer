import os

import time

import re

import random

from google.genai import Client, errors

from docx import Document

from docx.shared import Pt, Inches



# 1. Setup Gemini (Use your verified API Key)

client = Client(api_key="")



def get_ai_optimized_content(jd_text, master_prompt):

    """Uses stable model and handles both 429 (Quota) and 503 (Busy) errors."""

    full_query = (

        "You are an ATS Resume Expert. Return ONLY content for: SUMMARY:, SKILLS:, EXP_1:, EXP_2:, PROJECTS:.\n"

        "STRICT FORMATTING RULES:\n"

        "1. SUMMARY: Provide a single high-impact paragraph. NO bullets, NO '*' symbols.\n"

        "2. SKILLS/PROJECTS: Use **Category or Project Title** on its own line (NO bullet).\n"

        "3. Use '*' ONLY for the descriptive bullet points starting on the line BELOW a title.\n"

        "4. DO NOT use blank lines between a title and its bullets.\n"

        f"\n\nJD: {jd_text}\n\nPrompt: {master_prompt}"

    )

    

    MODEL_NAME = 'models/gemini-2.5-flash'

    

    for attempt in range(5): 

        try:

            print(f"Connecting to {MODEL_NAME} (Attempt {attempt + 1})...")

            response = client.models.generate_content(model=MODEL_NAME, contents=full_query)

            

            content = response.text

            data_map = {}

            sections = ["SUMMARY", "SKILLS", "EXP_1", "EXP_2", "PROJECTS"]

            for section in sections:

                # Optimized regex to capture text between headers

                pattern = rf"{section}:\s*(.*?)\s*(?={'|'.join(sections)}|$)"

                match = re.search(pattern, content, re.DOTALL)

                data_map[section] = match.group(1).strip() if match else ""

            return data_map



        except Exception as e:

            if "429" in str(e) or "503" in str(e):

                # Add jitter: (base wait) + (random 1-5 seconds)

                wait_time = ((attempt + 1) * 30) + random.randint(1, 5)

                print(f"Server busy or Quota hit. Retrying in {wait_time}s...")

                time.sleep(wait_time)

            else:

                print(f"Technical error: {e}")

                return None

    return None



def apply_formatted_text(paragraph, text, section_name):

    """

    Hardened formatting engine:

    1. Removes all bullet markers for SKILLS and SUMMARY.

    2. Correctly identifies and applies Word Bolding to text inside ** markers.

    3. Prevents stray asterisks from appearing in the final document.

    """

    font_name = "Times New Roman"

    font_size = Pt(11)

    

    if paragraph.runs:

        font_name = paragraph.runs[0].font.name or "Times New Roman"

        font_size = paragraph.runs[0].font.size or Pt(11)



    # Filter out empty lines to prevent lonely bullets

    lines = [l.strip() for l in text.split('\n') if l.strip()]

    current_p = paragraph

    

    for i, line_text in enumerate(lines):

        target_p = current_p if i == 0 else current_p.insert_paragraph_before("")

        if i > 0: current_p._element.addnext(target_p._element)

        

        # --- LOGIC: Handle Bullets ---

        is_bullet = line_text.startswith('*') or line_text.startswith('•')

        # Force Summary and Skills to stay flush-left (No Bullets)

        if section_name in ["SKILLS", "SUMMARY", "Keys"]:

            is_bullet = False



        if is_bullet:

            target_p.text = "•\t"

            target_p.paragraph_format.left_indent = Inches(0.25)

            target_p.paragraph_format.first_line_indent = Inches(-0.25)

            clean_line = line_text.lstrip('*• ').strip()

        else:

            target_p.text = ""

            target_p.paragraph_format.left_indent = Inches(0)

            target_p.paragraph_format.first_line_indent = Inches(0)

            clean_line = line_text



        # --- UPDATED BOLDING ENGINE ---

        # We split the line by bold markers but KEEP the markers in the split to identify them

        parts = re.split(r'(\*\*.*?\*\*)', clean_line)

        for part in parts:

            if not part: continue

            run = target_p.add_run()

            

            # Check if this specific part is a bolded section

            if part.startswith('**') and part.endswith('**'):

                # Strip the markers ONLY after confirming it is a bold section

                run.text = part.replace('**', '').strip()

                run.bold = True

            else:

                # Regular text

                run.text = part

            

            run.font.name = font_name

            run.font.size = font_size



        target_p.paragraph_format.space_after = Pt(2)

        target_p.paragraph_format.line_spacing = 1.0

        current_p = target_p



def update_resume(base_path, output_path, replacements):

    doc = Document(base_path)

    for p in list(doc.paragraphs):

        for tag_name, value in replacements.items():

            if "{{" + tag_name + "}}" in p.text:

                # Pass the tag_name so the engine knows the section-specific rules

                apply_formatted_text(p, value, tag_name)

    doc.save(output_path)

    print(f"\n--- SUCCESS! Resume tailored to: {output_path}")

# # --- DATA INPUTS ---
job_description = """ Job Summary

DESCRIPTION:

Duties: Participate in planning sessions with project managers, business analysts, and team members to analyze business requirements and outline the proposed IT solution to develop software applications. Participate in design reviews and provide input for design recommendations. Incorporate security requirements into design, and provide input to information and data flow for software application development. Understand and comply with Project Life Cycle Methodology in all planning steps and adhere to IT Control Policies throughout design, development, and testing. Incorporate Corporate Architectural Standards into application design specifications, document the detailed application specifications, and translate technical requirements into programmed application modules. Develop and enhance software application modules. Participate in programming reviews and ensure that all solutions are aligned with pre-defined architectural specifications. Identify and troubleshoot application programming-related issues and review and provide feedback on the final user documentation. Work with the test team to ensure the testing results correspond to the business expectations. Assist in project planning and tracking activities, and production support triaging activities.

QUALIFICATIONS:

Minimum education and experience required: Master's degree in Computer Science or related field of study plus 1 year (12 months) of experience in the job offered or as Software Engineer, Software Developer, or related occupation.

Skills Required: This position requires one (1) year of experience with the following: Develop and maintain full stack applications, ensuring robust enterprise infrastructure and adherence to cyber security controls; Collaborate with cross-functional teams to design and implement innovative solutions using Java, .NET/C#, Spring Boot, AWS, Apache Kafka, Elastic Search, and ReactJS; Utilize Core Java and the Servlet API for application development; Implement web services using SOAP and REST protocols; Design and develop asynchronous communication interfaces and messaging systems using Apache Kafka and UM (Universal Messaging) queues; Apply multithreading techniques, including thread pooling and synchronization, to enhance application performance; Conduct unit testing and adopt Test-Driven Development practices using Junit; Develop web UI components with HTML, CSS, and web debugging tools; Ensure efficient project delivery utilizing Agile and Scrum methodologies; Deploy applications on Apache Tomcat and manage messaging with Apache Kafka; Utilize relational databases such as Oracle and SQL Server, and apply Hibernate and PL/SQL for data management; Implement GraphQL for efficient data querying; Utilize Docker and Kubernetes to architect and implement a microservices-based system; Manage build processes using Maven, Git, and Jenkins; Integrate Continuous Integration/DevOps practices, including automated testing, build automation, and automated deployment, to streamline development workflows; Ensure application security through rigorous testing and validation using tools including SSAP Scan, SonarQube, and Splunk.
 
 """
master_prompt = """ Role: Act as an expert Technical Recruiter and ATS Optimization Specialist.

Task: Based on the Job Description (JD) provided below and my professional experience, rewrite my resume to be "Top 1%" caliber.

Strategic Instructions:

Infer the Ecosystem: Do not just match keywords. If the JD mentions a core language (e.g., Java, Python, JavaScript), automatically incorporate the industry-standard "hidden" dependencies (Frameworks, Testing libraries, CI/CD tools, and Cloud services) that a hiring manager would expect.

Quantified Impact: Rewrite my experience using the "Action Verb + Task + Measurable Result" formula. Ensure every bullet proves business value. Use percentages (%) for roughly 30% of the points, and use hard numbers (e.g., $2M, 50k users) for the rest.

The "Unique" Factor: Craft a high-impact Professional Summary as a single, strong paragraph. DO NOT use bullet points or the * symbol for this section.

Output Requirements:

Experience Section: Provide 10–11 high-impact bullet points for each role. Focus on transformational results, not just responsibilities.

Academic/Technical Projects: Create two (2) high-level project entries that align perfectly with the JD’s technical "intention." These should fill any gaps between my experience and the JD's requirements using the inferred tech stack.

Tech Stack Section: Categorize my skills into exactly 7 or fewer categories (e.g., Programming Languages, Frameworks, Cloud & DevOps). Format this as a simple, copiable list (e.g., "Category: x, y, z").

Seniority Check: Align the complexity of the tasks and projects strictly with the seniority level implied by the JD.

My Experience:

Full stack Developer, Capital One  - Plano, TX		June 2025 – PresentFull Stack Software Engineer, 
Full stack Developer, CGI - Benagluru, India 	May 2022 – July 2023
Additional Task: Rewrite my resume sections based on the JD.
IMPORTANT: Do not put a '' before Project Titles. A Project Title should look like: Project Name | Tech Stack. Use the '' ONLY for the bullet points underneath the project title.
CRITICAL RULES:
1. NO REPETITION: Do not include company names, roles, or dates in the output. Provide ONLY the bullet points or descriptions.
2. FORMATTING: Use **bold** for skill categories and project titles. 
3. EXPERIENCE: Provide exactly 10-11 bullet points per role. Start each bullet with a single '*'. 
4. PROJECTS: Provide 2 projects. Each should have a **Project Title** line followed by 3 bullet points starting with '*'. No '**1.' numbering.
5. TECH STACK: Use the format: **Category Name**: Skill 1, Skill 2..."""

try:
    ai_data = get_ai_optimized_content(job_description, master_prompt)
    
    if ai_data is not None:
        update_resume(
            base_path="C:/Users/abhin/OneDrive/Desktop/AbhinavKoritalaprac.docx",
            output_path="C:/Users/abhin/OneDrive/Desktop/AbhinavResume.docx",
            replacements=ai_data
        )
    else:
        print("FAILED: Gemini API was unavailable after multiple attempts.")

except Exception as e:
    print(f"CRITICAL ERROR: {e}")